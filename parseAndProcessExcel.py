from datetime import datetime, timedelta
import openpyxl
import sys


############################## Customize these values ###################################
VERSION = 'RFK/2.0'     # Official version tag
EXCEL_FILE = "DB LektionsPlaner.xlsx"

FIRST_LESSON=1  # Default=first lesson number, Change to start processing elsewhere
LAST_LESSON=3   # Default=last lesson number, chance to stop processing
FIRST_JOIN=FIRST_LESSON # Default=FIRST_LESSON, change this if you want files to be joined from here
LAST_JOIN =LAST_LESSON  # Default=last lesson number, Stop joining files here
REMOVE_WORD_FILES=True  # Default: True, remove temporary word files after converting
#########################################################################################

DEBUG=False
OVERVIEW_ROW_OFFSET=7 # Overview sheet has an row_offset of +7, so lesson 26 details are found at row # 33
DETAIL_ROW_OFFSET=5   # Detail data starts at row 5 on the sheet
MAX_ROWS_ON_FORM = 22 # Max number of rows available on the lesson plan



# Global variables
wb = openpyxl.load_workbook(EXCEL_FILE)
sheet=wb.active
context={'version':VERSION} 


def selectSheet(mySheet: str):
    """
    Changes the global variable "sheet" to a new sheet and gives a confirmation log

    Args:
        mySheet (str): Name of the new sheet, for example "Lesson 3"

    """
    globals()['sheet']=wb[mySheet]
    if DEBUG: print("Sheet selected " + sheet.title)


def getDuration(cellName : str) -> str:
    """
    Will check if the give cell contains a timeDate value, and if it does
    process the value and return a string with this duration.

    Args:
        cellName (str): Cell address, for example "A23"

    Returns:
        Duration formatted as "hh:mm" or an empty string

    """
    cell = sheet[cellName]

    if cell.data_type == 'd':
        cell_value = cell.value
        hours = cell_value.seconds // 3600
        minutes = (cell_value.seconds // 60) % 60
        return f"{hours}:{minutes:02d}"
    else:
        return ""

def getLessonDuration(lessonNumber: int):
    """
    Called from getLessonOverview and expect the global sheet to be set to "Overview"
    in advance of the call.

    Will check the cells in column C and D containing data for dual and solo duration 
    and return a string with a properly formatted text.

    Will return an empty string if there are no durations provided.

    Args:
        lessonNumber (int): Lesson number

    """
    lessonDuration=getDuration(f"C{lessonNumber+OVERVIEW_ROW_OFFSET}")

    if lessonDuration!= "":
        lessonDuration="Est. Dual: "+lessonDuration

    sololDuration= getDuration(f"D{lessonNumber+OVERVIEW_ROW_OFFSET}")
    if sololDuration!= "":
        if (lessonDuration==""):
            lessonDuration="Est. Solo: " + sololDuration
        else:
            lessonDuration=lessonDuration+", solo: "+sololDuration

    return lessonDuration


def getLessonOverview(lessonNumber : int):
    """
    Populate the global lookup table with the basic lesson details "name", "number" 
    and "duration". 
    
    Since these details are provided on the "overview" sheet, this sheet is globally 
    selected.

    Args:
        lessonNumber (int): Lesson number. 

    """
    selectSheet("Oversigt")

    globals()['context'].update({'lesson_number':sheet[f"A{lessonNumber+OVERVIEW_ROW_OFFSET}"].value})
    globals()['context'].update({'lesson_name':sheet[f"B{lessonNumber+OVERVIEW_ROW_OFFSET}"].value})

    myDurationString=getLessonDuration(lessonNumber)
    globals()['context'].update({'lesson_duration':myDurationString})

 

def numberOfLines():
    """
    Used internally by getLessonDetails to find number of lines for air work and for briefing topics.
    Check if there is room enough of the word template - if not, make a hard exit
    Otherwise return the two numbers to the caller.
    """
    noAirwork = 0
    noBriefing = 0

    # This is only called from getLessonDetails, so global sheet has already been set
    for row in sheet.iter_rows(min_row=5, max_row=sheet.max_row, values_only=True):
        if row[1] is not None and row != "":
            noAirwork += 1
        if row[3] is not None and row != "":
            noBriefing += 1
    
    ## Check if there is sufficient space for this lesson
    max_rows=MAX_ROWS_ON_FORM
    if noBriefing != 0:
        # We need a header, subtract one from the total number of available lines
        max_rows-=1

    if noBriefing + noAirwork > max_rows:
        print(f"Tabellerne i {sheet.title} indeholder mere end {MAX_ROWS_ON_FORM} rækker - det er der desværre ikke plads til i Word dokumentet.")
        sys.exit()
    else:
        return noAirwork, noBriefing


def getLessonDetails(lessonNumber: int):
    """
    Find the appropriate air exercises and briefing topics for the lesson on the lesson sheet.
    Then populate the global context dictionary with numbered bookmarks and these details.

    For briefing topics, start at the bottom and back upwards and add a bolded header before
    the briefing topics.
    
    Args:
        lessonNumber (int): Lesson number

    """

    import html  # used to escape the ampersands

    selectSheet(f"Lesson {lessonNumber}")

    noAirwork, noBriefing = numberOfLines()
    if DEBUG:
        print(f"Number of rows with airwork: {noAirwork}")
        print(f"Number of rows with briefings: {noBriefing}")
        print(" ")

    # Insert air exercises
    for row in range(DETAIL_ROW_OFFSET, DETAIL_ROW_OFFSET + noAirwork):
        value1 = float(sheet[f'A{row}'].value)
        value2 = sheet[f'B{row}'].value
    
        globals()['context'].update({f'N{row-DETAIL_ROW_OFFSET+1}':f'{value1:.2f}'})
        globals()['context'].update({f'T{row-DETAIL_ROW_OFFSET+1}':html.escape(value2)})

    import docx
    from docxtpl import RichText

    # Insert briefings
    if noBriefing > 0:
        # Insert briefing header
        bookmark_line = MAX_ROWS_ON_FORM - noBriefing
        bold_text = RichText()
        bold_text.add("Briefing topics:", bold=True)
        globals()['context'].update({f'T{bookmark_line}':bold_text})

        # Briefing contents
        for row in range(DETAIL_ROW_OFFSET, DETAIL_ROW_OFFSET + noBriefing):
            value1 = sheet[f'D{row}'].value
            value2 = sheet[f'E{row}'].value
        
            bookmark_line += 1
            globals()['context'].update({f'N{bookmark_line}':value1})
            globals()['context'].update({f'T{bookmark_line}':html.escape(value2)})



def generateGlobals(lessonNumber):
    getLessonOverview(lessonNumber)
    getLessonDetails(lessonNumber)

if DEBUG: 
    # myLessonNumber = 26
    # generateGlobals(myLessonNumber)
    print(globals()['context'])

####################################
### Patching subscripted speeds here
####################################

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_subscript(paragraph, main_text, subscript_text):
    """
    Adds main text followed by subscript text as new runs in the paragraph.
    """
    run = paragraph.add_run(main_text)
    sub_run = paragraph.add_run()
    t = OxmlElement('w:t')
    t.text = subscript_text
    sub_run._r.append(t)
    rPr = sub_run._element.get_or_add_rPr()
    sub = OxmlElement('w:vertAlign')
    sub.set(qn('w:val'), 'subscript')
    rPr.append(sub)
    if DEBUG: print(f"Added subscript: {main_text}{subscript_text}")

def contains_rich_text(paragraph):
    for run in paragraph.runs:
        # Check if the XML of the run contains rich text properties
        run_xml = run._element.xml
        if "<w:b" in run_xml or "<w:i" in run_xml or "<w:u" in run_xml or "<w:color" in run_xml or "<w:sz" in run_xml or "<w:rFonts" in run_xml:
            return True
    return False

def replace_text_with_subscript(paragraph, replacements):
    """
    Replaces all occurrences of the search_texts in a paragraph with subscript formatted subscript_texts.
    """

    # Avoid accidentally removing the bold text briefings header
    if contains_rich_text(paragraph): return  # Return immediately 

    # OK, we can proceed
    full_text = paragraph.text
    new_paragraph_text = []
    last_index = 0

    while last_index < len(full_text):
        found = False
        for search_text in replacements:
            start_idx = full_text.find(search_text, last_index)
            if start_idx != -1:
                end_idx = start_idx + len(search_text)
                if end_idx<len(full_text) and full_text[end_idx] in " ,.?" or end_idx==len(full_text):
                    if start_idx > 0 : new_paragraph_text.append(full_text[last_index:start_idx])
                    new_paragraph_text.append((search_text[0], search_text[1:]))
                    last_index = end_idx
                    found = True
            
        if not found:
            new_paragraph_text.append(full_text[last_index:])
            break

    if DEBUG: print("Constructed new_paragraph_text:", new_paragraph_text)

    # Clear the paragraph runs
    for run in paragraph.runs:
        run.text = ""

    # Add the reconstructed text
    for item in new_paragraph_text:
        if isinstance(item, tuple):
            main_text, subscript_text = item
            add_subscript(paragraph, main_text, subscript_text)
        else:
            paragraph.add_run(item)
    if DEBUG: print(f"Processed paragraph: {paragraph.text}")

def process_paragraph(paragraph, replacements):
    """
    Processes a single paragraph to replace text with subscript.
    """
    if DEBUG: print(f"Original paragraph: {paragraph.text}")
    replace_text_with_subscript(paragraph, replacements)

def process_table(table, replacements):
    """
    Processes all cells in a table to replace text with subscript.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph, replacements)

def patch_subscripts(doc):
    # Define the replacements: (search_text, subscript_text)
    # Necessary because subscripts are lost in tranlation from excel to word :-(
    replacements = [ "VX", "VY", "VA", "VR", "VS", "VS0", "VS1", "VNO", "VNE", "VFE", "VREF", "VGLIDE" ]

    # Replace the terms with subscripts in tables
    for table in doc.tables:
        process_table(table, replacements)




####################################
### Looping over lessons starts here
####################################

from docxtpl import DocxTemplate, RichText
import docx2pdf

for lesson in range(FIRST_LESSON, LAST_LESSON+1):
    generateGlobals(lesson)
    doc=DocxTemplate("LektionsSkabelon.docx")
    doc.render(context)

    # WORK IN PROGRESS....
    patch_subscripts(doc)

    tempName=f"temp_lesson_{lesson}"
    doc.save(tempName+".docx")

    # Reset document template again with a fresh globals
    context={'version':VERSION}

    # Convert the Word document to PDF
    docx2pdf.convert(tempName+".docx", tempName+".pdf")

# Join all pdf lesson documents from start to finish and create "lesson-plan-[version].pdf"
import os
if LAST_JOIN-FIRST_JOIN > 1:
    # List comprehension is a good way to generate names
    lessons = " ".join([f"temp_lesson_{i}.pdf" for i in range(FIRST_JOIN, LAST_JOIN+1)])
    cmd = "pdfunite "+lessons+ f" lesson-plan-{context['version'].split('/')[1]}.pdf"

    # Use a system command to activate "pdfunite" with all this
    os.system(cmd)

    print(f"Document: lesson-plan-{context['version'].split('/')[1]}.pdf ready\n\n")

if REMOVE_WORD_FILES:
    import glob
    print("Cleaning up...")
    for file in glob.glob("temp_lesson_*.docx"):
        os.remove(file)
    print("Temporary word files have been successfully removed")
else:
    print("Temporary word files remain in current working directory")


# #####
# Move pdf lesson document pdf files to backup version folder 
import re,shutil

# Create a subfolder for this version files - if not already present
subfolder_name = f"version_{context['version']}"
os.makedirs(subfolder_name, exist_ok=True)

# We need a few paths prepared
current_dir = os.getcwd()
subfolder_path=os.path.join(current_dir, subfolder_name)

# And a regular expression pattern to identify file names 
pattern = r"temp_lesson_\d+\.pdf"

# Now loop through matching files in the current directory
for filename in os.listdir(current_dir):
    if re.match(pattern, filename):
        # Construct the source and destination paths
        source_path = os.path.join(current_dir, filename)
        dest_path = os.path.join(subfolder_path, filename)

        # Move file to the new subfolder
        shutil.move(source_path, dest_path)
        if DEBUG: print(f"Moved {filename} to {subfolder_name}")

print(f"The individual lesson plans has been successfully backed up to: {subfolder_path}")